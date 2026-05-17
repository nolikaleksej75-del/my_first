import telebot
from telebot.types import InlineKeyboardMarkup, InlineKeyboardButton
import json
import os
from datetime import datetime, timedelta
import traceback
import time
import threading

# ============ КОНФИГУРАЦИЯ ============
TOKEN = '8717752281:AAFcE_fNSLQHoFgEY8EYV0dVnp5wlqwZGpU'
ADMIN_ID = 1531814351

# Проверка токена
if TOKEN == 'YOUR_BOT_TOKEN_HERE':
    print("=" * 60)
    print("❌ ОШИБКА: Токен бота не настроен!")
    print("=" * 60)
    exit(1)

# Импорт для создания DOCX
try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    import io
    DOCX_AVAILABLE = True
    print("✅ Библиотека python-docx загружена")
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ Установите python-docx: pip install python-docx")

bot = telebot.TeleBot(TOKEN)

try:
    bot.get_me()
    print("✅ Подключение к Telegram API успешно!")
except Exception as e:
    print(f"❌ Ошибка: {e}")
    exit(1)

# ============ ДАННЫЕ ============
DATA_FILE = 'zapis.json'
SETTINGS_FILE = 'settings.json'
COMMANDERS_FILE = 'commanders.json'
USERS_FILE = 'users.json'
BLOCKED_USERS_FILE = 'blocked_users.json'
CUSTOM_NAMES_FILE = 'custom_names.json'

# Группы
GROUPS = {
    '721': '721',
    '722': '722',
    '723': '723',
    '724a': '724а',
    '724b': '724б',
    '725': '725',
    '726': '726',
    '727a': '727а',
    '727b': '727б'
}

# Дни недели
DAYS_RU = {
    'monday': 'Понедельник',
    'tuesday': 'Вторник',
    'wednesday': 'Среда',
    'thursday': 'Четверг',
    'friday': 'Пятница',
    'saturday': 'Суббота',
    'sunday': 'Воскресенье'
}

DAY_NAMES_TO_KEY = {v: k for k, v in DAYS_RU.items()}

# Структура для хранения командиров и старшины
DATA_STRUCTURE = {
    'commanders': {
        721: [], 722: [], 723: [], 724: [], 725: [], 726: [], 727: []
    },
    'starhina': []
}

# ============ ФУНКЦИИ ДЛЯ РАБОТЫ С ДАННЫМИ ============
def load_users():
    if os.path.exists(USERS_FILE):
        try:
            with open(USERS_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except:
            return []
    return []

def save_users(users):
    try:
        with open(USERS_FILE, 'w', encoding='utf-8') as f:
            json.dump(users, f, ensure_ascii=False, indent=2)
    except:
        pass

def add_user(user_id, user_info):
    users = load_users()
    for user in users:
        if user['id'] == user_id:
            return
    users.append({
        'id': user_id,
        'first_name': user_info.first_name,
        'username': user_info.username,
        'last_name': user_info.last_name,
        'added_at': datetime.now().strftime("%d.%m.%Y %H:%M:%S")
    })
    save_users(users)

def load_blocked_users():
    if os.path.exists(BLOCKED_USERS_FILE):
        try:
            with open(BLOCKED_USERS_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except:
            return []
    return []

def save_blocked_users(blocked):
    try:
        with open(BLOCKED_USERS_FILE, 'w', encoding='utf-8') as f:
            json.dump(blocked, f, ensure_ascii=False, indent=2)
    except:
        pass

def is_user_blocked(user_id):
    return user_id in load_blocked_users()

def load_custom_names():
    if os.path.exists(CUSTOM_NAMES_FILE):
        try:
            with open(CUSTOM_NAMES_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except:
            return {}
    return {}

def save_custom_names(names):
    try:
        with open(CUSTOM_NAMES_FILE, 'w', encoding='utf-8') as f:
            json.dump(names, f, ensure_ascii=False, indent=2)
    except:
        pass

def get_display_name(user_id):
    custom_names = load_custom_names()
    if str(user_id) in custom_names:
        return custom_names[str(user_id)]
    users = load_users()
    for user in users:
        if user['id'] == user_id:
            if user.get('username'):
                return f"@{user['username']}"
            return user.get('first_name', str(user_id))
    return str(user_id)

def get_user_mention(user_id):
    custom_names = load_custom_names()
    if str(user_id) in custom_names:
        return custom_names[str(user_id)]
    users = load_users()
    for user in users:
        if user['id'] == user_id:
            if user.get('username'):
                return f"@{user['username']}"
            return user.get('first_name', str(user_id))
    return str(user_id)

def load_commanders_data():
    global DATA_STRUCTURE
    if os.path.exists(COMMANDERS_FILE):
        try:
            with open(COMMANDERS_FILE, 'r', encoding='utf-8') as f:
                loaded = json.load(f)
                if isinstance(loaded, dict):
                    if 'commanders' in loaded and 'starhina' in loaded:
                        DATA_STRUCTURE = loaded
                    else:
                        DATA_STRUCTURE = {'commanders': {721: [], 722: [], 723: [], 724: [], 725: [], 726: [], 727: []}, 'starhina': []}
                        for key, value in loaded.items():
                            if str(key).isdigit() and int(key) in DATA_STRUCTURE['commanders']:
                                if isinstance(value, list):
                                    DATA_STRUCTURE['commanders'][int(key)] = value
                save_commanders_data()
        except Exception as e:
            print(f"Ошибка загрузки командиров: {e}")
            DATA_STRUCTURE = {'commanders': {721: [], 722: [], 723: [], 724: [], 725: [], 726: [], 727: []}, 'starhina': []}
            save_commanders_data()
    else:
        save_commanders_data()

def save_commanders_data():
    try:
        with open(COMMANDERS_FILE, 'w', encoding='utf-8') as f:
            json.dump(DATA_STRUCTURE, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"Ошибка сохранения командиров: {e}")

def is_admin(user_id):
    return user_id == ADMIN_ID

def is_starhina(user_id):
    return user_id in DATA_STRUCTURE.get('starhina', [])

def is_commander(user_id):
    commanders = DATA_STRUCTURE.get('commanders', {})
    for platoon, lst in commanders.items():
        if user_id in lst:
            return platoon
    return None

def has_commander_permissions(user_id):
    if is_admin(user_id):
        return True
    if is_starhina(user_id):
        return True
    if is_commander(user_id) is not None:
        return True
    return False

def get_user_role_string(user_id):
    if is_admin(user_id):
        return "👑 Главный администратор"
    if is_starhina(user_id):
        return "⭐ Старшина курса"
    platoon = is_commander(user_id)
    if platoon:
        group_names = {721: '721', 722: '722', 723: '723', 724: '724', 725: '725', 726: '726', 727: '727'}
        return f"🎖️ Командир {group_names.get(platoon, str(platoon))} группы"
    return "👤 Пользователь"

def get_available_platoons_for_user(user_id):
    if is_admin(user_id):
        return list(DAYS_RU.keys())
    if is_starhina(user_id):
        return list(DAYS_RU.keys())
    platoon = is_commander(user_id)
    if platoon:
        return [platoon]
    return []

def get_group_number_from_platoon(platoon):
    group_map = {721: '721', 722: '722', 723: '723', 724: '724', 725: '725', 726: '726', 727: '727'}
    return group_map.get(platoon, str(platoon))

def load_data():
    if os.path.exists(DATA_FILE):
        try:
            with open(DATA_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except:
            return {day: [] for day in DAYS_RU.keys()}
    return {day: [] for day in DAYS_RU.keys()}

def save_data(data):
    try:
        with open(DATA_FILE, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except:
        pass

def load_settings():
    if os.path.exists(SETTINGS_FILE):
        try:
            with open(SETTINGS_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except:
            return {day: True for day in DAYS_RU.keys()}
    return {day: True for day in DAYS_RU.keys()}

def save_settings(settings):
    try:
        with open(SETTINGS_FILE, 'w', encoding='utf-8') as f:
            json.dump(settings, f, ensure_ascii=False, indent=2)
    except:
        pass

def is_day_enabled(day_key):
    settings = load_settings()
    return settings.get(day_key, True)

def get_next_date_for_day(day_key):
    today = datetime.now()
    days_map = {'monday': 0, 'tuesday': 1, 'wednesday': 2, 'thursday': 3, 'friday': 4, 'saturday': 5, 'sunday': 6}
    target_weekday = days_map[day_key]
    days_ahead = target_weekday - today.weekday()
    if days_ahead <= 0:
        days_ahead += 7
    return today + timedelta(days=days_ahead)

def get_russian_day_name(date):
    days = ['Понедельник', 'Вторник', 'Среда', 'Четверг', 'Пятница', 'Суббота', 'Воскресенье']
    return days[date.weekday()]

def notify_commander_about_new_record(platoon, record_info):
    commanders = DATA_STRUCTURE.get('commanders', {}).get(platoon, [])
    starhinas = DATA_STRUCTURE.get('starhina', [])
    all_to_notify = commanders + starhinas
    if not all_to_notify:
        return
    message_text = (
        f"🔔 **НОВАЯ ЗАПИСЬ В ВАШЕЙ ГРУППЕ!** 🔔\n\n"
        f"👤 **Курсант:** {record_info['surname']}\n"
        f"⭐ **Звание:** {record_info.get('rank', 'рядовой')}\n"
        f"📅 **День:** {record_info['day_name']}\n"
        f"🕐 **Время записи:** {record_info['time']}\n"
        f"🆔 **ID курсанта:** `{record_info['user_id']}`\n\n"
        f"💡 /commander – управление записями"
    )
    for cid in all_to_notify:
        try:
            bot.send_message(cid, message_text, parse_mode='Markdown')
            time.sleep(0.05)
        except:
            pass

# ============ ФУНКЦИЯ СОЗДАНИЯ СПИСКА ============
def create_duty_list_for_day(target_date, records):
    if not DOCX_AVAILABLE:
        return None
    
    doc = Document()
    
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)
    
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(15)
    style.paragraph_format.line_spacing = Pt(14)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    
    # ШАПКА
    approve_indent = Cm(9.0)
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = approve_indent
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("УТВЕРЖДАЮ")
    run.bold = False
    run.font.size = Pt(15)
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = approve_indent
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Начальник авиационного")
    run.font.size = Pt(15)
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = approve_indent
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("факультета")
    run.font.size = Pt(15)
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = approve_indent
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    spaces_15 = " " * 15
    run = p.add_run(f"полковник{spaces_15}О.Л.Турчинович")
    run.font.size = Pt(15)
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = approve_indent
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f" .{target_date.strftime('%m')}.{target_date.strftime('%Y')}")
    run.font.size = Pt(15)
    
    doc.add_paragraph()
    
    # ЗАГОЛОВОК "СПИСОК"
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("СПИСОК")
    run.bold = False
    run.font.size = Pt(15)
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("личного состава 2 курса ")
    run.font.size = Pt(15)
    run.add_break()
    run = p.add_run("авиационного факультета, ")
    run.font.size = Pt(15)
    run.add_break()
    run = p.add_run("убывающего в увольнение")
    run.font.size = Pt(15)
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(target_date.strftime("%d.%m.%Y"))
    run.font.size = Pt(15)
    
    doc.add_paragraph()
    
    # ТАБЛИЦА
    headers = ['№\nп/п', 'Воинское\nзвание', 'Фамилия\nи инициалы', 'Учебная\nгруппа', 
               'До которого\nчисла и часа\nубыл', 'Какого числа\nи в котором\nчасу вернулся', 
               'Роспись лица,\nпринявшего\nдоклад о\nприбытии']
    
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    
    table.columns[0].width = Cm(1.0)
    table.columns[1].width = Cm(3.26)
    table.columns[2].width = Cm(4.48)
    table.columns[3].width = Cm(1.5)
    table.columns[4].width = Cm(2.5)
    table.columns[5].width = Cm(2.25)
    table.columns[6].width = Cm(2.5)
    
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.bold = False
            run.font.size = Pt(12)
        cell.paragraphs[0].line_spacing = Pt(12)
    
    row = table.add_row()
    cell = row.cells[0]
    cell.merge(row.cells[6])
    cell.text = "Личный состав убывающий с 18:00"
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cell.paragraphs[0].runs:
        run.font.bold = False
        run.font.size = Pt(12)
    cell.paragraphs[0].line_spacing = Pt(12)
    
    sorted_records = sorted(records, key=lambda x: (x.get('platoon', 0), x.get('surname', '')))
    departure_time = f"{target_date.strftime('%d.%m.%Y')} 23:00"
    
    for idx, record in enumerate(sorted_records, 1):
        row = table.add_row()
        
        cell = row.cells[0]
        cell.text = f"{idx}."
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(12)
        cell.paragraphs[0].line_spacing = Pt(12)
        
        cell = row.cells[1]
        rank = record.get('rank', 'рядовой')
        cell.text = rank
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(12)
        cell.paragraphs[0].line_spacing = Pt(12)
        
        cell = row.cells[2]
        surname = record.get('surname', '')
        cell.text = surname
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(12)
        cell.paragraphs[0].line_spacing = Pt(12)
        
        cell = row.cells[3]
        platoon = record.get('platoon', 0)
        group_num = get_group_number_from_platoon(platoon)
        cell.text = group_num
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(12)
        cell.paragraphs[0].line_spacing = Pt(12)
        
        cell = row.cells[4]
        cell.text = departure_time
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(12)
        cell.paragraphs[0].line_spacing = Pt(12)
        
        cell = row.cells[5]
        cell.text = "            "
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        cell = row.cells[6]
        cell.text = "            "
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    doc.add_paragraph()
    
    # ПОДПИСИ
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Начальник 2 курса")
    run.font.size = Pt(15)
    p.paragraph_format.line_spacing = Pt(14)
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("батальона курсантов")
    run.font.size = Pt(15)
    p.paragraph_format.line_spacing = Pt(14)
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("авиационного факультета")
    run.font.size = Pt(15)
    p.paragraph_format.line_spacing = Pt(14)
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    spaces_85 = " " * 85
    run = p.add_run(f"капитан{spaces_85}А.К.Садовский")
    run.font.size = Pt(15)
    p.paragraph_format.line_spacing = Pt(14)
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"   .{target_date.strftime('%m')}.{target_date.strftime('%Y')}")
    run.font.size = Pt(15)
    p.paragraph_format.line_spacing = Pt(14)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Исполняющий обязанности")
    run.font.size = Pt(15)
    p.paragraph_format.line_spacing = Pt(14)
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("командира батальона курсантов")
    run.font.size = Pt(15)
    p.paragraph_format.line_spacing = Pt(14)
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("авиационного факультета")
    run.font.size = Pt(15)
    p.paragraph_format.line_spacing = Pt(14)
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    spaces_85 = " " * 85
    run = p.add_run(f"майор{spaces_85}П.Е.Русакович")
    run.font.size = Pt(15)
    p.paragraph_format.line_spacing = Pt(14)
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"   .{target_date.strftime('%m')}.{target_date.strftime('%Y')}")
    run.font.size = Pt(15)
    p.paragraph_format.line_spacing = Pt(14)
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return file_stream

def generate_and_send_list(target_date, chat_id):
    if not DOCX_AVAILABLE:
        bot.send_message(chat_id, "❌ Библиотека python-docx не установлена!")
        return False, "Библиотека не установлена"
    
    data = load_data()
    
    day_name_ru = get_russian_day_name(target_date)
    day_key = None
    for key, name in DAYS_RU.items():
        if name == day_name_ru:
            day_key = key
            break
    
    if not day_key:
        return False, f"Не удалось определить день недели"
    
    records = data.get(day_key, [])
    
    if not records:
        return False, f"Нет записей на {day_name_ru}"
    
    try:
        file_stream = create_duty_list_for_day(target_date, records)
        if file_stream:
            filename = f"Список_увольняемых_{target_date.strftime('%d.%m.%Y')}.docx"
            bot.send_document(chat_id, (filename, file_stream), 
                            caption=f"📋 Список увольняемых на {day_name_ru} {target_date.strftime('%d.%m.%Y')}")
            return True, f"Список на {day_name_ru} успешно создан"
        else:
            return False, "Ошибка при создании документа"
    except Exception as e:
        return False, f"Ошибка: {str(e)}"

# ============ АВТОМАТИЧЕСКОЕ СОЗДАНИЕ ПО РАСПИСАНИЮ ============
def check_and_generate_scheduled_lists():
    while True:
        now = datetime.now()
        
        if now.hour == 21 and now.minute == 30:
            weekday = now.weekday()
            
            if weekday == 2:
                target_date = now + timedelta(days=1)
                success, msg = generate_and_send_list(target_date, ADMIN_ID)
                print(f"[{now}] Среда -> Четверг: {msg}")
                time.sleep(60)
            
            elif weekday == 3:
                for days_ahead in [1, 2, 3]:
                    target_date = now + timedelta(days=days_ahead)
                    success, msg = generate_and_send_list(target_date, ADMIN_ID)
                    print(f"[{now}] Четверг -> {get_russian_day_name(target_date)}: {msg}")
                    time.sleep(5)
                time.sleep(60)
        
        time.sleep(30)

# ============ ФУНКЦИИ РАССЫЛКИ ============
def send_broadcast(message_text, sender_id):
    users = load_users()
    success_count = 0
    fail_count = 0
    
    for user in users:
        try:
            bot.send_message(user['id'], message_text, parse_mode='Markdown')
            success_count += 1
            time.sleep(0.05)
        except Exception as e:
            fail_count += 1
    
    report = f"✅ **Рассылка завершена!**\n\n📨 **Отправлено:** {success_count}\n❌ **Не доставлено:** {fail_count}\n📊 **Всего:** {len(users)}"
    bot.send_message(sender_id, report, parse_mode='Markdown')
    return success_count, fail_count

# ============ КЛАВИАТУРЫ ============
def main_menu_keyboard():
    kb = InlineKeyboardMarkup(row_width=2)
    kb.add(
        InlineKeyboardButton("📝 Записаться", callback_data="new_zapis"),
        InlineKeyboardButton("📋 Мои записи", callback_data="my_records"),
        InlineKeyboardButton("ℹ️ О боте", callback_data="about_bot"),
        InlineKeyboardButton("🆔 Мой ID", callback_data="show_my_id")
    )
    return kb

def days_keyboard():
    kb = InlineKeyboardMarkup(row_width=2)
    btns = []
    for dk, dn in DAYS_RU.items():
        if is_day_enabled(dk):
            btns.append(InlineKeyboardButton(f"✅ {dn}", callback_data=f"day_{dk}"))
        else:
            btns.append(InlineKeyboardButton(f"❌ {dn}", callback_data="disabled_day"))
    for i in range(0, len(btns), 2):
        if i+1 < len(btns):
            kb.add(btns[i], btns[i+1])
        else:
            kb.add(btns[i])
    kb.add(InlineKeyboardButton("🏠 Главное меню", callback_data="main_menu"))
    return kb

def group_keyboard():
    kb = InlineKeyboardMarkup(row_width=3)
    buttons = []
    for group_key, group_name in GROUPS.items():
        buttons.append(InlineKeyboardButton(group_name, callback_data=f"group_{group_key}"))
    for i in range(0, len(buttons), 3):
        if i+2 < len(buttons):
            kb.add(buttons[i], buttons[i+1], buttons[i+2])
        elif i+1 < len(buttons):
            kb.add(buttons[i], buttons[i+1])
        else:
            kb.add(buttons[i])
    kb.add(InlineKeyboardButton("🔙 Назад к дням", callback_data="back_to_days"))
    kb.add(InlineKeyboardButton("🏠 Главное меню", callback_data="main_menu"))
    return kb

def commander_keyboard():
    kb = InlineKeyboardMarkup(row_width=1)
    kb.add(
        InlineKeyboardButton("📋 Список записей", callback_data="commander_list"),
        InlineKeyboardButton("❌ Удалить запись", callback_data="commander_remove")
    )
    kb.add(InlineKeyboardButton("🏠 Главное меню", callback_data="main_menu"))
    return kb

def admin_keyboard():
    kb = InlineKeyboardMarkup(row_width=1)
    kb.add(
        InlineKeyboardButton("📅 Управление днями", callback_data="admin_days"),
        InlineKeyboardButton("📊 Все записи", callback_data="admin_list"),
        InlineKeyboardButton("🗑 Очистить записи", callback_data="admin_clear"),
        InlineKeyboardButton("🗑 Очистить день", callback_data="admin_clear_day"),
        InlineKeyboardButton("👥 Управление командирами", callback_data="admin_commanders"),
        InlineKeyboardButton("⭐ Управление старшиной", callback_data="admin_starhina"),
        InlineKeyboardButton("🔒 Блокировка пользователей", callback_data="admin_block_users"),
        InlineKeyboardButton("✏️ Установить имя", callback_data="admin_set_name"),
        InlineKeyboardButton("🔔 Закрыть запись", callback_data="admin_close_day"),
        InlineKeyboardButton("📢 Сделать рассылку", callback_data="admin_broadcast"),
        InlineKeyboardButton("📋 Создать список", callback_data="admin_create_list"),
        InlineKeyboardButton("🔄 Создать все списки", callback_data="admin_create_all_lists")
    )
    kb.add(InlineKeyboardButton("🏠 Главное меню", callback_data="main_menu"))
    return kb

def admin_days_keyboard():
    kb = InlineKeyboardMarkup(row_width=2)
    sets = load_settings()
    for dk, dn in DAYS_RU.items():
        status = "✅" if sets[dk] else "❌"
        kb.add(InlineKeyboardButton(f"{status} {dn}", callback_data=f"toggle_{dk}"))
    kb.add(InlineKeyboardButton("🔙 Назад", callback_data="admin_back"))
    return kb

def admin_clear_day_keyboard():
    kb = InlineKeyboardMarkup(row_width=2)
    for dk, dn in DAYS_RU.items():
        kb.add(InlineKeyboardButton(dn, callback_data=f"clear_day_{dk}"))
    kb.add(InlineKeyboardButton("🔙 Назад", callback_data="admin_back"))
    return kb

def admin_close_day_keyboard():
    kb = InlineKeyboardMarkup(row_width=2)
    for dk, dn in DAYS_RU.items():
        kb.add(InlineKeyboardButton(dn, callback_data=f"close_day_{dk}"))
    kb.add(InlineKeyboardButton("🔙 Назад", callback_data="admin_back"))
    return kb

def admin_commanders_keyboard():
    kb = InlineKeyboardMarkup(row_width=1)
    group_names = ['721', '722', '723', '724', '725', '726', '727']
    for grp in group_names:
        cnt = len(DATA_STRUCTURE['commanders'].get(int(grp), []))
        kb.add(InlineKeyboardButton(f"📌 {grp} группа ({cnt} ком.)", callback_data=f"manage_platoon_{grp}"))
    kb.add(InlineKeyboardButton("🔙 Назад", callback_data="admin_back"))
    return kb

def manage_platoon_keyboard(platoon):
    kb = InlineKeyboardMarkup(row_width=1)
    kb.add(InlineKeyboardButton("➕ Добавить командира", callback_data=f"add_commander_{platoon}"))
    for cid in DATA_STRUCTURE['commanders'].get(platoon, []):
        name = get_display_name(cid)
        kb.add(InlineKeyboardButton(f"❌ Удалить {name}", callback_data=f"remove_commander_{platoon}_{cid}"))
    kb.add(InlineKeyboardButton("🔙 Назад", callback_data="admin_commanders"))
    return kb

def admin_starhina_keyboard():
    kb = InlineKeyboardMarkup(row_width=1)
    kb.add(InlineKeyboardButton("➕ Добавить старшину", callback_data="add_starhina"))
    for sid in DATA_STRUCTURE.get('starhina', []):
        name = get_display_name(sid)
        kb.add(InlineKeyboardButton(f"❌ Удалить {name}", callback_data=f"remove_starhina_{sid}"))
    kb.add(InlineKeyboardButton("🔙 Назад", callback_data="admin_back"))
    return kb

def admin_block_users_keyboard():
    kb = InlineKeyboardMarkup(row_width=1)
    kb.add(InlineKeyboardButton("➕ Заблокировать", callback_data="block_user"))
    kb.add(InlineKeyboardButton("🔓 Разблокировать", callback_data="unblock_user"))
    blocked = load_blocked_users()
    if blocked:
        kb.add(InlineKeyboardButton("📋 Список заблокированных", callback_data="list_blocked"))
    kb.add(InlineKeyboardButton("🔙 Назад", callback_data="admin_back"))
    return kb

def confirm_keyboard():
    kb = InlineKeyboardMarkup(row_width=2)
    kb.add(
        InlineKeyboardButton("✅ ДА", callback_data="confirm_clear"),
        InlineKeyboardButton("❌ НЕТ", callback_data="admin_back")
    )
    return kb

def confirm_clear_day_keyboard(day_key):
    kb = InlineKeyboardMarkup(row_width=2)
    kb.add(
        InlineKeyboardButton("✅ ДА, очистить", callback_data=f"confirm_clear_day_{day_key}"),
        InlineKeyboardButton("❌ НЕТ", callback_data="admin_back")
    )
    return kb

def my_records_keyboard(records):
    kb = InlineKeyboardMarkup(row_width=1)
    for i, rec in enumerate(records):
        kb.add(InlineKeyboardButton(f"❌ {rec['day_name']} - {rec['surname']}", callback_data=f"delete_my_record_{i}"))
    kb.add(InlineKeyboardButton("🔙 Назад", callback_data="main_menu"))
    return kb

# ============ ПЕРЕМЕННЫЕ ДЛЯ ВВОДА ============
waiting_for_rank = {}
waiting_for_surname = {}
broadcast_waiting = {}

# ============ КОМАНДЫ ============
@bot.message_handler(commands=['start'])
def start_command(message):
    try:
        if is_user_blocked(message.from_user.id):
            bot.send_message(message.chat.id, "⛔ Вы заблокированы!")
            return
        add_user(message.from_user.id, message.from_user)
        role = get_user_role_string(message.from_user.id)
        text = f"🌟 **Добро пожаловать!** 🌟\n\n⭐ **Ваша роль:** {role}\n\n📌 **Как записаться:**\n1️⃣ Выберите день\n2️⃣ Выберите группу\n3️⃣ Выберите звание\n4️⃣ Введите фамилию с инициалами\n\n🎖️ **Командиры:** /commander\n👑 **Админ:** /admin\n🆔 **Мой ID:** /myid"
        bot.send_message(message.chat.id, text, parse_mode='Markdown', reply_markup=main_menu_keyboard())
    except Exception as e:
        print(f"Ошибка: {e}")

@bot.message_handler(commands=['myid'])
def myid_command(message):
    role = get_user_role_string(message.from_user.id)
    text = f"🆔 **Ваш ID:** `{message.from_user.id}`\n⭐ **Роль:** {role}"
    bot.send_message(message.chat.id, text, parse_mode='Markdown')

@bot.message_handler(commands=['admin'])
def admin_command(message):
    if not is_admin(message.from_user.id):
        bot.send_message(message.chat.id, "⛔ Доступ запрещен!")
        return
    bot.send_message(message.chat.id, "👑 **Панель администратора**", parse_mode='Markdown', reply_markup=admin_keyboard())

@bot.message_handler(commands=['commander'])
def commander_command(message):
    if not has_commander_permissions(message.from_user.id):
        bot.send_message(message.chat.id, "⛔ Вы не командир и не старшина.")
        return
    role = get_user_role_string(message.from_user.id)
    bot.send_message(message.chat.id, f"🎖️ **Панель {role}**", parse_mode='Markdown', reply_markup=commander_keyboard())

@bot.message_handler(func=lambda message: True)
def handle_text_input(message):
    user_id = message.from_user.id
    
    # Обработка рассылки
    if broadcast_waiting.get(user_id):
        msg_text = message.text
        broadcast_waiting[user_id] = False
        confirm_kb = InlineKeyboardMarkup(row_width=2)
        confirm_kb.add(
            InlineKeyboardButton("✅ ДА, ОТПРАВИТЬ", callback_data="confirm_broadcast"),
            InlineKeyboardButton("❌ ОТМЕНА", callback_data="admin_back")
        )
        broadcast_waiting['pending'] = {'user_id': user_id, 'text': msg_text}
        bot.send_message(message.chat.id, f"📢 **Предпросмотр:**\n\n{msg_text}\n\nОтправить?", 
                        parse_mode='Markdown', reply_markup=confirm_kb)
        return
    
    # Обработка ручного ввода звания
    if user_id in waiting_for_rank:
        rank = message.text.strip().lower()
        if not rank:
            bot.send_message(message.chat.id, "❌ Введите звание:")
            return
        
        temp_data = waiting_for_rank[user_id]
        temp_data['rank'] = rank
        waiting_for_surname[user_id] = temp_data
        del waiting_for_rank[user_id]
        
        bot.send_message(message.chat.id, f"✅ Звание: {rank}\n\n✏️ **Введите фамилию с инициалами**\n\nПример: `Иванов И.И.`",
                        parse_mode='Markdown')
        return
    
    # Обработка фамилии
    if user_id in waiting_for_surname:
        surname = message.text.strip()
        if not surname:
            bot.send_message(message.chat.id, "❌ Введите фамилию с инициалами:")
            return
        
        temp_data = waiting_for_surname[user_id]
        day_key = temp_data['day_key']
        day_name = temp_data['day_name']
        group_key = temp_data['group_key']
        rank = temp_data.get('rank', 'рядовой')
        platoon = int(group_key.replace('a', '').replace('b', ''))
        
        if not is_day_enabled(day_key):
            bot.send_message(message.chat.id, f"❌ День {day_name} закрыт")
            del waiting_for_surname[user_id]
            return
        
        all_data = load_data()
        record = {
            'surname': surname,
            'rank': rank,
            'user_id': user_id,
            'username': message.from_user.username or '',
            'first_name': message.from_user.first_name or '',
            'platoon': platoon,
            'group': GROUPS[group_key],
            'day_name': day_name,
            'time': datetime.now().strftime("%d.%m.%Y %H:%M:%S")
        }
        all_data[day_key].append(record)
        save_data(all_data)
        
        notify_commander_about_new_record(platoon, record)
        
        del waiting_for_surname[user_id]
        
        temp_file = f'temp_{user_id}.json'
        if os.path.exists(temp_file):
            os.remove(temp_file)
        
        bot.send_message(message.chat.id, f"✅ **Запись оформлена!**\n📅 {day_name}\n👤 {surname}\n⭐ {rank}\n📌 {GROUPS[group_key]} группа\n🕐 {record['time']}",
                        parse_mode='Markdown', reply_markup=main_menu_keyboard())

# ============ CALLBACK ============
@bot.callback_query_handler(func=lambda call: True)
def callback_handler(call):
    try:
        if call.data == "main_menu":
            role = get_user_role_string(call.from_user.id)
            bot.edit_message_text(f"🌟 **Главное меню**\n\n⭐ **Ваша роль:** {role}",
                                 call.message.chat.id, call.message.message_id,
                                 parse_mode='Markdown', reply_markup=main_menu_keyboard())
        
        elif call.data == "about_bot":
            bot.edit_message_text("ℹ️ **О боте**\n\n🤖 Версия 5.0\n📅 Запись по дням и группам\n⭐ Выбор воинского звания\n📋 Автоматическое создание списков\n⏰ Время убытия: 23:00\n\n🎖️ Командиры и старшина могут управлять записями",
                                 call.message.chat.id, call.message.message_id,
                                 parse_mode='Markdown', reply_markup=main_menu_keyboard())
        
        elif call.data == "show_my_id":
            bot.answer_callback_query(call.id)
            bot.send_message(call.message.chat.id, f"🆔 Ваш ID: `{call.from_user.id}`", parse_mode='Markdown')
        
        elif call.data == "my_records":
            data = load_data()
            my_records = []
            for dk, dn in DAYS_RU.items():
                for rec in data[dk]:
                    if isinstance(rec, dict) and rec.get('user_id') == call.from_user.id:
                        my_records.append({'day_key': dk, 'day_name': dn, 'surname': rec['surname'], 
                                          'time': rec['time'], 'group': rec.get('group'), 'rank': rec.get('rank', '')})
            if not my_records:
                bot.answer_callback_query(call.id, "У вас нет активных записей", show_alert=True)
                return
            with open(f'my_records_{call.from_user.id}.json', 'w', encoding='utf-8') as f:
                json.dump(my_records, f)
            text = "📋 **Ваши записи:**\n\n"
            for i, rec in enumerate(my_records):
                text += f"{i+1}. {rec['day_name']} - {rec['surname']} ({rec.get('rank', '')}, {rec.get('group', '?')})\n   🕐 {rec['time']}\n\n"
            bot.send_message(call.message.chat.id, text, parse_mode='Markdown', reply_markup=my_records_keyboard(my_records))
        
        elif call.data.startswith('delete_my_record_'):
            idx = int(call.data.split('_')[-1])
            with open(f'my_records_{call.from_user.id}.json', 'r', encoding='utf-8') as f:
                my_records = json.load(f)
            if idx < len(my_records):
                record = my_records[idx]
                data = load_data()
                for dk in DAYS_RU.keys():
                    for i, rec in enumerate(data[dk]):
                        if (rec.get('user_id') == call.from_user.id and 
                            rec.get('surname') == record['surname'] and
                            rec.get('time') == record['time']):
                            data[dk].pop(i)
                            save_data(data)
                            break
                if os.path.exists(f'my_records_{call.from_user.id}.json'):
                    os.remove(f'my_records_{call.from_user.id}.json')
                bot.answer_callback_query(call.id, "✅ Запись удалена!")
                bot.send_message(call.message.chat.id, "✅ Запись удалена!", reply_markup=main_menu_keyboard())
            else:
                bot.answer_callback_query(call.id, "❌ Ошибка")
        
        elif call.data == "new_zapis":
            if is_user_blocked(call.from_user.id):
                bot.answer_callback_query(call.id, "⛔ Вы заблокированы!", show_alert=True)
                return
            bot.edit_message_text("📅 **Выберите день недели:**", call.message.chat.id, call.message.message_id,
                                 parse_mode='Markdown', reply_markup=days_keyboard())
        
        elif call.data == "back_to_days":
            bot.edit_message_text("📅 **Выберите день недели:**", call.message.chat.id, call.message.message_id,
                                 parse_mode='Markdown', reply_markup=days_keyboard())
        
        elif call.data.startswith('day_') and call.data != "disabled_day":
            if is_user_blocked(call.from_user.id):
                bot.answer_callback_query(call.id, "⛔ Вы заблокированы!", show_alert=True)
                return
            day_key = call.data.replace('day_', '')
            if not is_day_enabled(day_key):
                bot.answer_callback_query(call.id, "❌ День закрыт", show_alert=True)
                return
            day_name = DAYS_RU[day_key]
            with open(f'temp_{call.from_user.id}.json', 'w', encoding='utf-8') as f:
                json.dump({'day_key': day_key, 'day_name': day_name}, f)
            bot.edit_message_text(f"✅ Вы выбрали: {day_name}\n\n📌 **Выберите группу:**",
                                 call.message.chat.id, call.message.message_id,
                                 parse_mode='Markdown', reply_markup=group_keyboard())
        
        elif call.data.startswith('group_'):
            if is_user_blocked(call.from_user.id):
                bot.answer_callback_query(call.id, "⛔ Вы заблокированы!", show_alert=True)
                return
            group_key = call.data.replace('group_', '')
            temp_file = f'temp_{call.from_user.id}.json'
            if not os.path.exists(temp_file):
                bot.answer_callback_query(call.id, "❌ Ошибка", show_alert=True)
                return
            with open(temp_file, 'r', encoding='utf-8') as f:
                tmp = json.load(f)
            
            waiting_for_rank[call.from_user.id] = {
                'day_key': tmp['day_key'],
                'day_name': tmp['day_name'],
                'group_key': group_key
            }
            
            rank_kb = InlineKeyboardMarkup(row_width=1)
            rank_kb.add(
                InlineKeyboardButton("⭐ рядовой", callback_data="rank_рядовой"),
                InlineKeyboardButton("⭐ младший сержант", callback_data="rank_младший сержант"),
                InlineKeyboardButton("⭐ сержант", callback_data="rank_сержант")
            )
            rank_kb.add(InlineKeyboardButton("✏️ Ввести свое", callback_data="rank_custom"))
            
            bot.edit_message_text(f"✅ Вы выбрали: {GROUPS[group_key]} группа\n\n⭐ **Выберите воинское звание:**",
                                 call.message.chat.id, call.message.message_id,
                                 parse_mode='Markdown', reply_markup=rank_kb)
            bot.answer_callback_query(call.id)
        
        elif call.data.startswith('rank_'):
            rank_value = call.data.replace('rank_', '')
            
            if call.from_user.id not in waiting_for_rank:
                bot.answer_callback_query(call.id, "❌ Ошибка! Начните запись заново", show_alert=True)
                return
            
            temp_data = waiting_for_rank[call.from_user.id]
            
            if rank_value == 'custom':
                bot.edit_message_text(f"✅ Вы выбрали: {GROUPS[temp_data['group_key']]} группа\n\n✏️ **Введите ваше звание**\n\nПримеры: ефрейтор, старший сержант, старшина",
                                     call.message.chat.id, call.message.message_id,
                                     parse_mode='Markdown')
                bot.answer_callback_query(call.id)
                return
            else:
                temp_data['rank'] = rank_value
                waiting_for_surname[call.from_user.id] = temp_data
                del waiting_for_rank[call.from_user.id]
                
                bot.edit_message_text(f"✅ Вы выбрали: {GROUPS[temp_data['group_key']]} группа\n⭐ Звание: {rank_value}\n\n✏️ **Введите фамилию с инициалами**\n\nПример: `Иванов И.И.`",
                                     call.message.chat.id, call.message.message_id,
                                     parse_mode='Markdown')
                bot.answer_callback_query(call.id)
        
        elif call.data == "disabled_day":
            bot.answer_callback_query(call.id, "❌ День закрыт для записи", show_alert=True)
        
        # ============ КОМАНДИР ============
        elif call.data == "commander_list":
            if not has_commander_permissions(call.from_user.id):
                bot.answer_callback_query(call.id, "⛔ Нет прав")
                return
            data = load_data()
            text = "📋 ЗАПИСИ:\n\n"
            found = False
            for dk, dn in DAYS_RU.items():
                if data[dk]:
                    for rec in data[dk]:
                        if isinstance(rec, dict):
                            display_name = get_display_name(rec.get('user_id', 0))
                            text += f"📌 {dn}: {rec.get('surname', '?')} ({rec.get('rank', '')}) - {rec.get('group', '?')}\n   ID: {rec.get('user_id')}\n   Время: {rec.get('time', '')}\n\n"
                            found = True
            if not found:
                text += "Нет записей"
            bot.send_message(call.message.chat.id, text, reply_markup=commander_keyboard())
            bot.answer_callback_query(call.id)
        
        elif call.data == "commander_remove":
            if not has_commander_permissions(call.from_user.id):
                bot.answer_callback_query(call.id, "⛔ Нет прав")
                return
            data = load_data()
            records = []
            for dk, dn in DAYS_RU.items():
                for idx, rec in enumerate(data[dk]):
                    if isinstance(rec, dict):
                        records.append({'day_key': dk, 'day_name': dn, 'index': idx, 'record': rec})
            if not records:
                bot.answer_callback_query(call.id, "Нет записей для удаления", show_alert=True)
                return
            with open(f'commander_records_{call.from_user.id}.json', 'w', encoding='utf-8') as f:
                json.dump(records, f)
            kb = InlineKeyboardMarkup(row_width=1)
            for i, r in enumerate(records):
                display_name = get_display_name(r['record']['user_id'])
                kb.add(InlineKeyboardButton(f"{r['day_name']} - {r['record']['surname']} ({display_name})", callback_data=f"remove_select_{i}"))
            kb.add(InlineKeyboardButton("🔙 Назад", callback_data="commander_back"))
            bot.edit_message_text("❌ **Выберите запись для удаления:**", call.message.chat.id, call.message.message_id, reply_markup=kb)
        
        elif call.data.startswith('remove_select_'):
            if not has_commander_permissions(call.from_user.id):
                bot.answer_callback_query(call.id, "⛔ Нет прав")
                return
            idx = int(call.data.split('_')[-1])
            with open(f'commander_records_{call.from_user.id}.json', 'r', encoding='utf-8') as f:
                recs = json.load(f)
            sel = recs[idx]
            with open(f'remove_selected_{call.from_user.id}.json', 'w', encoding='utf-8') as f:
                json.dump(sel, f)
            deleter_info = get_user_role_string(call.from_user.id)
            bot.edit_message_text(f"❌ **Удаление записи**\n\n👤 {sel['record']['surname']}\n📅 {sel['day_name']}\n\n✏️ **Введите причину удаления:**",
                                 call.message.chat.id, call.message.message_id)
            bot.register_next_step_handler(call.message, lambda m: process_remove_reason(m, deleter_info))
        
        elif call.data == "commander_back":
            role = get_user_role_string(call.from_user.id)
            bot.edit_message_text(f"🎖️ **Панель {role}**", call.message.chat.id, call.message.message_id,
                                 parse_mode='Markdown', reply_markup=commander_keyboard())
        
        # ============ АДМИН ============
        elif call.data == "admin_days":
            if not is_admin(call.from_user.id):
                return
            bot.edit_message_text("📅 **Управление днями**", call.message.chat.id, call.message.message_id,
                                 reply_markup=admin_days_keyboard())
        
        elif call.data.startswith('toggle_'):
            if not is_admin(call.from_user.id):
                return
            dk = call.data.replace('toggle_', '')
            sets = load_settings()
            sets[dk] = not sets[dk]
            save_settings(sets)
            bot.answer_callback_query(call.id, f"День {DAYS_RU[dk]} {'включен' if sets[dk] else 'отключен'}")
            bot.edit_message_text("📅 **Управление днями**", call.message.chat.id, call.message.message_id,
                                 reply_markup=admin_days_keyboard())
        
        elif call.data == "admin_list":
            if not is_admin(call.from_user.id):
                return
            data = load_data()
            text = "📊 ВСЕ ЗАПИСИ:\n\n"
            total = 0
            for dk, dn in DAYS_RU.items():
                if data[dk]:
                    text += f"📌 {dn}:\n"
                    for rec in data[dk]:
                        if isinstance(rec, dict):
                            text += f"   • {rec.get('surname', '?')} ({rec.get('rank', '')}) - {rec.get('group', '?')}\n     🕐 {rec.get('time', '')}\n"
                            total += 1
                    text += "\n"
            if total == 0:
                text += "Нет записей"
            bot.send_message(call.message.chat.id, text)
            bot.answer_callback_query(call.id)
        
        elif call.data == "admin_clear":
            if not is_admin(call.from_user.id):
                return
            bot.edit_message_text("⚠️ **Очистить все записи?**", call.message.chat.id, call.message.message_id,
                                 reply_markup=confirm_keyboard())
        
        elif call.data == "confirm_clear":
            if not is_admin(call.from_user.id):
                return
            save_data({day: [] for day in DAYS_RU.keys()})
            bot.answer_callback_query(call.id, "✅ Все записи очищены!")
            bot.edit_message_text("👑 **Панель администратора**", call.message.chat.id, call.message.message_id,
                                 parse_mode='Markdown', reply_markup=admin_keyboard())
        
        elif call.data == "admin_clear_day":
            if not is_admin(call.from_user.id):
                return
            bot.edit_message_text("🗑 **Очистить записи на день**", call.message.chat.id, call.message.message_id,
                                 reply_markup=admin_clear_day_keyboard())
        
        elif call.data.startswith("clear_day_"):
            if not is_admin(call.from_user.id):
                return
            dk = call.data.replace('clear_day_', '')
            bot.edit_message_text(f"⚠️ **Очистить записи на {DAYS_RU[dk]}?**", call.message.chat.id, call.message.message_id,
                                 reply_markup=confirm_clear_day_keyboard(dk))
        
        elif call.data.startswith("confirm_clear_day_"):
            if not is_admin(call.from_user.id):
                return
            dk = call.data.replace('confirm_clear_day_', '')
            data = load_data()
            data[dk] = []
            save_data(data)
            bot.answer_callback_query(call.id, f"✅ Записи на {DAYS_RU[dk]} очищены!")
            bot.edit_message_text("👑 **Панель администратора**", call.message.chat.id, call.message.message_id,
                                 parse_mode='Markdown', reply_markup=admin_keyboard())
        
        elif call.data == "admin_commanders":
            if not is_admin(call.from_user.id):
                return
            bot.edit_message_text("👥 **Управление командирами**", call.message.chat.id, call.message.message_id,
                                 reply_markup=admin_commanders_keyboard())
        
        elif call.data.startswith("manage_platoon_"):
            if not is_admin(call.from_user.id):
                return
            pl = int(call.data.split('_')[-1])
            bot.edit_message_text(f"👥 **Управление {pl} группой**", call.message.chat.id, call.message.message_id,
                                 reply_markup=manage_platoon_keyboard(pl))
        
        elif call.data.startswith("add_commander_"):
            if not is_admin(call.from_user.id):
                return
            pl = int(call.data.split('_')[-1])
            bot.answer_callback_query(call.id)
            msg = bot.send_message(call.message.chat.id, f"➕ **Добавить командира для {pl} группы**\n\nОтправьте ID пользователя:\n\nПример: `1531814351`", parse_mode='Markdown')
            bot.register_next_step_handler(msg, lambda m: add_commander(m, pl, call.message))
        
        elif call.data.startswith("remove_commander_"):
            if not is_admin(call.from_user.id):
                return
            parts = call.data.split('_')
            pl = int(parts[2])
            cid = int(parts[3])
            if cid in DATA_STRUCTURE['commanders'][pl]:
                DATA_STRUCTURE['commanders'][pl].remove(cid)
                save_commanders_data()
                bot.answer_callback_query(call.id, "✅ Командир удален")
                bot.edit_message_text(f"👥 **Управление {pl} группой**", call.message.chat.id, call.message.message_id,
                                     reply_markup=manage_platoon_keyboard(pl))
        
        elif call.data == "admin_starhina":
            if not is_admin(call.from_user.id):
                return
            bot.edit_message_text("⭐ **Управление старшиной**", call.message.chat.id, call.message.message_id,
                                 reply_markup=admin_starhina_keyboard())
        
        elif call.data == "add_starhina":
            if not is_admin(call.from_user.id):
                return
            bot.answer_callback_query(call.id)
            msg = bot.send_message(call.message.chat.id, "⭐ **Добавить старшину**\n\nОтправьте ID пользователя:\n\nПример: `1531814351`", parse_mode='Markdown')
            bot.register_next_step_handler(msg, lambda m: add_starhina(m, call.message))
        
        elif call.data.startswith("remove_starhina_"):
            if not is_admin(call.from_user.id):
                return
            sid = int(call.data.split('_')[-1])
            if sid in DATA_STRUCTURE['starhina']:
                DATA_STRUCTURE['starhina'].remove(sid)
                save_commanders_data()
                bot.answer_callback_query(call.id, "✅ Старшина удален")
                bot.edit_message_text("⭐ **Управление старшиной**", call.message.chat.id, call.message.message_id,
                                     reply_markup=admin_starhina_keyboard())
        
        elif call.data == "admin_block_users":
            if not is_admin(call.from_user.id):
                return
            bot.edit_message_text("🔒 **Управление блокировкой**", call.message.chat.id, call.message.message_id,
                                 reply_markup=admin_block_users_keyboard())
        
        elif call.data == "block_user":
            if not is_admin(call.from_user.id):
                return
            bot.answer_callback_query(call.id)
            msg = bot.send_message(call.message.chat.id, "🔒 **Заблокировать пользователя**\n\nОтправьте ID:\n\nПример: `1531814351`", parse_mode='Markdown')
            bot.register_next_step_handler(msg, lambda m: block_user_by_id(m, 'block'))
        
        elif call.data == "unblock_user":
            if not is_admin(call.from_user.id):
                return
            bot.answer_callback_query(call.id)
            msg = bot.send_message(call.message.chat.id, "🔓 **Разблокировать пользователя**\n\nОтправьте ID:\n\nПример: `1531814351`", parse_mode='Markdown')
            bot.register_next_step_handler(msg, lambda m: block_user_by_id(m, 'unblock'))
        
        elif call.data == "list_blocked":
            if not is_admin(call.from_user.id):
                return
            blocked = load_blocked_users()
            if not blocked:
                text = "📋 Список заблокированных пуст"
            else:
                text = "📋 **Заблокированные:**\n\n"
                for uid in blocked:
                    text += f"• {get_display_name(uid)} (ID: `{uid}`)\n"
            bot.send_message(call.message.chat.id, text, parse_mode='Markdown', reply_markup=admin_block_users_keyboard())
        
        elif call.data == "admin_set_name":
            if not is_admin(call.from_user.id):
                return
            bot.answer_callback_query(call.id)
            msg = bot.send_message(call.message.chat.id, "✏️ **Установить имя курсанту**\n\nОтправьте ID и имя:\n`ID Имя`\n\nПример: `1531814351 Иванов И.И.`", parse_mode='Markdown')
            bot.register_next_step_handler(msg, set_custom_name)
        
        elif call.data == "admin_close_day":
            if not is_admin(call.from_user.id):
                return
            bot.edit_message_text("🔔 **Закрыть запись на день**", call.message.chat.id, call.message.message_id,
                                 reply_markup=admin_close_day_keyboard())
        
        elif call.data.startswith("close_day_"):
            if not is_admin(call.from_user.id):
                return
            dk = call.data.replace('close_day_', '')
            day_name = DAYS_RU[dk]
            users = load_users()
            sent = 0
            for user in users:
                try:
                    bot.send_message(user['id'], f"🔔 Запись на {day_name} закрыта! Списки будут готовы через 30 минут.")
                    sent += 1
                    time.sleep(0.05)
                except:
                    pass
            bot.answer_callback_query(call.id, f"Оповещение отправлено {sent} пользователям")
            bot.edit_message_text("👑 **Панель администратора**", call.message.chat.id, call.message.message_id,
                                 parse_mode='Markdown', reply_markup=admin_keyboard())
        
        elif call.data == "admin_broadcast":
            if not is_admin(call.from_user.id):
                return
            broadcast_waiting[call.from_user.id] = True
            bot.send_message(call.message.chat.id, "📢 **Режим рассылки**\n\nОтправьте текст сообщения:")
            bot.answer_callback_query(call.id)
        
        elif call.data == "confirm_broadcast":
            if not is_admin(call.from_user.id):
                return
            if broadcast_waiting.get('pending'):
                msg_text = broadcast_waiting['pending']['text']
                user_id = broadcast_waiting['pending']['user_id']
                bot.edit_message_text("📢 **Начинаю рассылку...**", call.message.chat.id, call.message.message_id)
                send_broadcast(msg_text, user_id)
                broadcast_waiting['pending'] = None
            bot.edit_message_text("👑 **Панель администратора**", call.message.chat.id, call.message.message_id,
                                 parse_mode='Markdown', reply_markup=admin_keyboard())
        
        elif call.data == "admin_create_list":
            if not is_admin(call.from_user.id):
                return
            if not DOCX_AVAILABLE:
                bot.answer_callback_query(call.id, "❌ Установите python-docx!", show_alert=True)
                return
            kb = InlineKeyboardMarkup(row_width=2)
            for dk, dn in DAYS_RU.items():
                kb.add(InlineKeyboardButton(dn, callback_data=f"create_list_{dk}"))
            kb.add(InlineKeyboardButton("🔙 Назад", callback_data="admin_back"))
            bot.edit_message_text("📋 **Выберите день:**", call.message.chat.id, call.message.message_id, reply_markup=kb)
        
        elif call.data.startswith("create_list_"):
            if not is_admin(call.from_user.id):
                return
            day_key = call.data.replace('create_list_', '')
            target_date = get_next_date_for_day(day_key)
            bot.send_message(call.message.chat.id, f"📋 **Создаю список на {DAYS_RU[day_key]}...**")
            success, msg = generate_and_send_list(target_date, call.message.chat.id)
            bot.answer_callback_query(call.id, msg)
        
        elif call.data == "admin_create_all_lists":
            if not is_admin(call.from_user.id):
                return
            if not DOCX_AVAILABLE:
                bot.answer_callback_query(call.id, "❌ Установите python-docx!", show_alert=True)
                return
            bot.send_message(call.message.chat.id, "📋 **Создаю списки на все дни...**")
            data = load_data()
            created = 0
            for day_key, records in data.items():
                if records:
                    target_date = get_next_date_for_day(day_key)
                    success, msg = generate_and_send_list(target_date, call.message.chat.id)
                    if success:
                        created += 1
                    time.sleep(2)
            bot.send_message(call.message.chat.id, f"✅ Создано списков: {created}")
            bot.answer_callback_query(call.id, f"Создано {created} списков")
        
        elif call.data == "admin_back":
            if not is_admin(call.from_user.id):
                return
            bot.edit_message_text("👑 **Панель администратора**", call.message.chat.id, call.message.message_id,
                                 parse_mode='Markdown', reply_markup=admin_keyboard())
        
        else:
            bot.answer_callback_query(call.id, "Неизвестная команда")
    
    except Exception as e:
        print(f"Ошибка: {e}")
        traceback.print_exc()

# ============ ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ ============
def process_remove_reason(message, deleter_info):
    try:
        reason = message.text.strip()
        if not reason:
            reason = "Не указана"
        with open(f'remove_selected_{message.from_user.id}.json', 'r', encoding='utf-8') as f:
            sel = json.load(f)
        day_key = sel['day_key']
        idx = sel['index']
        record = sel['record']
        data = load_data()
        if day_key in data and idx < len(data[day_key]):
            data[day_key].pop(idx)
            save_data(data)
            try:
                bot.send_message(record['user_id'], f"⚠️ **Ваша запись удалена!**\n📅 {sel['day_name']}\n👤 {record['surname']}\n❌ Причина: {reason}\n\n🗑 **Удалил:** {deleter_info}", parse_mode='Markdown')
            except:
                pass
            bot.send_message(message.chat.id, f"✅ Запись удалена!\nПричина: {reason}", reply_markup=commander_keyboard())
        for f in [f'commander_records_{message.from_user.id}.json', f'remove_selected_{message.from_user.id}.json']:
            if os.path.exists(f):
                os.remove(f)
    except Exception as e:
        print(f"Ошибка удаления: {e}")
        bot.send_message(message.chat.id, "❌ Ошибка при удалении")

def add_commander(msg, platoon, orig_msg):
    try:
        if not is_admin(msg.from_user.id):
            return
        user_input = msg.text.strip()
        if ' ' in user_input:
            user_input = user_input.split()[0]
        if user_input.startswith('@'):
            user_input = user_input[1:]
        if not user_input.isdigit():
            bot.send_message(msg.chat.id, f"❌ Ошибка! Нужно отправить ТОЛЬКО ID пользователя.\n\n✅ Правильно: `1531814351`", parse_mode='Markdown')
            return
        commander_id = int(user_input)
        if platoon not in DATA_STRUCTURE['commanders']:
            DATA_STRUCTURE['commanders'][platoon] = []
        if commander_id not in DATA_STRUCTURE['commanders'][platoon]:
            DATA_STRUCTURE['commanders'][platoon].append(commander_id)
            save_commanders_data()
            try:
                bot.send_message(commander_id, f"🎖️ **Поздравляем!**\n\nВы назначены командиром **{platoon} группы**!\n\n/commander", parse_mode='Markdown')
            except:
                pass
            bot.send_message(msg.chat.id, f"✅ Командир добавлен! ID: `{commander_id}`", parse_mode='Markdown')
        else:
            bot.send_message(msg.chat.id, f"❌ Уже является командиром!", parse_mode='Markdown')
        bot.send_message(orig_msg.chat.id, f"👥 **Управление {platoon} группой**", parse_mode='Markdown', reply_markup=manage_platoon_keyboard(platoon))
    except Exception as e:
        print(f"Ошибка: {e}")

def add_starhina(msg, orig_msg):
    try:
        if not is_admin(msg.from_user.id):
            return
        user_input = msg.text.strip()
        if ' ' in user_input:
            user_input = user_input.split()[0]
        if user_input.startswith('@'):
            user_input = user_input[1:]
        if not user_input.isdigit():
            bot.send_message(msg.chat.id, f"❌ Ошибка! Нужно отправить ТОЛЬКО ID пользователя.\n\n✅ Правильно: `1531814351`", parse_mode='Markdown')
            return
        sid = int(user_input)
        if sid not in DATA_STRUCTURE['starhina']:
            DATA_STRUCTURE['starhina'].append(sid)
            save_commanders_data()
            try:
                bot.send_message(sid, f"⭐ **Поздравляем!**\n\nВы назначены **Старшиной курса**!\n\n/commander", parse_mode='Markdown')
            except:
                pass
            bot.send_message(msg.chat.id, f"✅ Старшина добавлен! ID: `{sid}`", parse_mode='Markdown')
        else:
            bot.send_message(msg.chat.id, f"❌ Уже является старшиной!", parse_mode='Markdown')
        bot.send_message(orig_msg.chat.id, "⭐ **Управление старшиной**", parse_mode='Markdown', reply_markup=admin_starhina_keyboard())
    except Exception as e:
        print(f"Ошибка: {e}")

def block_user_by_id(msg, action_type):
    try:
        if not is_admin(msg.from_user.id):
            return
        user_input = msg.text.strip()
        if ' ' in user_input:
            user_input = user_input.split()[0]
        if user_input.startswith('@'):
            user_input = user_input[1:]
        if not user_input.isdigit():
            bot.send_message(msg.chat.id, f"❌ Ошибка! Нужно отправить ТОЛЬКО ID пользователя.", parse_mode='Markdown')
            return
        target_id = int(user_input)
        blocked = load_blocked_users()
        if action_type == 'block':
            if target_id not in blocked:
                blocked.append(target_id)
                save_blocked_users(blocked)
                bot.send_message(msg.chat.id, f"✅ Пользователь `{target_id}` заблокирован!", parse_mode='Markdown')
        else:
            if target_id in blocked:
                blocked.remove(target_id)
                save_blocked_users(blocked)
                bot.send_message(msg.chat.id, f"✅ Пользователь `{target_id}` разблокирован!", parse_mode='Markdown')
        bot.send_message(msg.chat.id, "🔒 **Управление блокировками**", parse_mode='Markdown', reply_markup=admin_block_users_keyboard())
    except Exception as e:
        print(f"Ошибка: {e}")

def set_custom_name(msg):
    try:
        if not is_admin(msg.from_user.id):
            return
        parts = msg.text.strip().split(maxsplit=1)
        if len(parts) != 2:
            bot.send_message(msg.chat.id, "❌ Формат: `ID Имя`\n\nПример: `1531814351 Иванов И.И.`", parse_mode='Markdown')
            return
        user_id = parts[0].strip()
        if not user_id.isdigit():
            bot.send_message(msg.chat.id, "❌ ID должен быть числом!", parse_mode='Markdown')
            return
        custom_name = parts[1].strip()
        custom_names = load_custom_names()
        custom_names[user_id] = custom_name
        save_custom_names(custom_names)
        bot.send_message(msg.chat.id, f"✅ Пользователю `{user_id}` установлено имя: **{custom_name}**", parse_mode='Markdown')
        bot.send_message(msg.chat.id, "👑 **Панель администратора**", parse_mode='Markdown', reply_markup=admin_keyboard())
    except Exception as e:
        print(f"Ошибка: {e}")

# ============ ЗАПУСК ============
if __name__ == '__main__':
    try:
        print("=" * 50)
        print("🤖 ЗАПУСК БОТА")
        print("=" * 50)
        print(f"👑 Админ ID: {ADMIN_ID}")
        print(f"📁 Папка: {os.getcwd()}")
        print("=" * 50)
        
        for file, default in [(SETTINGS_FILE, {day: True for day in DAYS_RU.keys()}), 
                              (DATA_FILE, {day: [] for day in DAYS_RU.keys()}),
                              (USERS_FILE, []), (BLOCKED_USERS_FILE, []), (CUSTOM_NAMES_FILE, {})]:
            if not os.path.exists(file):
                if isinstance(default, dict):
                    save_settings(default)
                elif isinstance(default, list):
                    save_users(default) if file == USERS_FILE else save_blocked_users(default)
                else:
                    save_custom_names(default)
                print(f"✅ Создан {file}")
        
        load_commanders_data()
        
        scheduler = threading.Thread(target=check_and_generate_scheduled_lists, daemon=True)
        scheduler.start()
        print("✅ Планировщик запущен (21:30 каждый день)")
        
        print("\n✅ БОТ ГОТОВ!")
        print("📢 Напишите /start в Telegram\n")
        print("=" * 50)
        
        bot.infinity_polling(timeout=60)
    
    except KeyboardInterrupt:
        print("\n⚠️ Бот остановлен")
    except Exception as e:
        print(f"❌ Ошибка: {e}")
        traceback.print_exc()
